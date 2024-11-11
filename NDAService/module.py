from PIL import Image 
from pytesseract import pytesseract 
from docx import Document
from PIL import Image, ImageOps
from aiogram.types import InputMediaPhoto, Message
from aiogram.fsm.storage.base import StorageKey
from aiogram.fsm.context import FSMContext
from config import checker_ids
import os
from config import PATHS
from PyPDF2 import PdfReader
from aiogram.types import FSInputFile


def GetDocx(endFilename, data):
    file_name = f"{PATHS['data']}nda.docx"
    document = Document(file_name)

    fio = f"{data.get('Фамилия').capitalize()} {data.get('Имя').capitalize()} {data.get('Отчество').capitalize()}"
    document.paragraphs[4].text = f'{fio}, именуемый/ая в дальнейшем как Получающая Сторона, и ТОО “Stratton”, именуемое в дальнейшем (Передающая сторона), в лице Директора Абдиханова Адила Алмасовича, действующего на основании Устава, с другой стороны, заключили настоящее Соглашение о неразглашении конфиденциальной информации (далее – Соглашение) о нижеследующем.'
    document.tables[-1].columns[-1].cells[1].text = fio
    document.tables[-1].columns[-1].cells[3].text = f"ИИН: {data.get('ИИН')} \nEmail:"
    document.save(endFilename)


def Upscale(fileName):
    img_orig = Image.open(fileName)

    img = ImageOps.scale(img_orig, 6, resample=Image.LANCZOS)
    img.save(fileName)


def ReadPhoto(fileName, mainSide=True):
    Upscale(fileName)
    path_to_tesseract = PATHS['tesseract']

    img = Image.open(fileName) 
    pytesseract.tesseract_cmd = path_to_tesseract 
    text = pytesseract.image_to_string(img, lang='rus+kaz').split()

    fields = {'ИИН': None, 
              'Имя': None, 
              'Фамилия': None, 
              'Отчество': None, 
              'Дата рождения': None} if mainSide else {'Дата выдачи': None}
    try:
        if mainSide:
            for i in range(len(text)):
                if text[i].lower().find('имя')!=-1:
                    fields['Имя']=text[i+1].capitalize()
                elif text[i].lower().find('фамилия')!=-1:
                    fields['Фамилия']=text[i+1].capitalize()
                elif text[i].lower().find('отчество')!=-1:
                    fields['Отчество']=text[i+1].capitalize()
                elif text[i].lower().find('рождения')!=-1:
                    fields['Дата рождения']=text[i+1]
                elif text[i].isdigit() and len(text[i])==12:
                    fields['ИИН']=text[i]
        else:
            for i in range(len(text)):
                if text[i].lower().find('действия')!=-1:
                    fields['Дата выдачи']=text[i+1]
                    break
    except:
        return 0
    finally:
        os.remove(fileName)
    return fields


def ReadPDF(fileName):
    try:
        fields = dict()
        reader = PdfReader(fileName)
        text = reader.pages[0].extract_text().split('\n')
        
        if len(text[0].split())==1:
            fields['Фамилия'] = text[0]
        if len(text[1].split())==1:
            fields['Имя'] = text[1]
        if len(text[2].split())==1:
            fields['Отчество'] = text[2]
        if len(text[3].split('.'))==3:
            fields['Дата рождения'] = text[3]
        if len(text[4])==12:
            fields['ИИН'] = text[4]
        return fields
    except Exception as ex:
        print(ex)
        return None
    finally:
        pass


def ReadDOCX(file_name):
    document = Document(file_name)

    for i in document.tables[-1].columns[-1].cells:
        print(i.text)


def ReadTextRequsites(text):
    fields = dict()
    text = text.split('\n')
    
    for i in text:
        if i.find('ИИК')>=0:
            fields['ИИК'] = ''.join(i.split()).split('ИИК')[-1]
        elif i.find('БИН')>=0:
            fields['БИН'] = ''.join(i.split()).split('БИН')[-1]
        # elif text.index(i)==0:
        #     fields['Компания'] = i
        elif text.index(i)==len(text)-1:
            fields['Лицо'] = i
    return fields


def GetMonoText(text):
    nv = ''
    for i in text:
        if i in [',', '.']:
            nv+='\\'+i
        else:
            nv+=i
    return nv


def RemoveMonoText(text):
    nv = ''
    for i in text:
        if i in ['\\', '`']:
            continue
        else:
            nv+=i
    return nv


async def SendToAdmin(msg, state, keyboard, fileName=None, requisites=None, document=None):
    data = await state.get_data()

    if fileName:
        for i in checker_ids:
            print(fileName)
            await msg.bot.send_document(chat_id=i, document=FSInputFile(fileName))
    elif requisites:
        for i in checker_ids:
            await msg.bot.send_message(chat_id=i, text=requisites)
    else:
        media = []
        photoPaths = [data['fImageId'], data['sImageId']]
        for path in photoPaths:
            media.append(InputMediaPhoto(media=path))

        for i in checker_ids:
            await msg.bot.send_media_group(chat_id=i, media=media)

    text1 = f'Пользователь @{msg.from_user.username} отправил данные вам на проверку для подписания {document}: '

    
    text2 = ' \n'.join([f'{k}\-{f"`{GetMonoText(str(v))}`" if v else "нет данных"}' for k, v in data.items() if k not in['fImageId','sImageId','face','back', 'pdf']])

    for i in checker_ids:
        storageKey = StorageKey(msg.bot.id, i, i)
        opponentState = FSMContext(storage=state.storage, key=storageKey)
        asd = {'text':f'{text1}\n{text2}'}
        if fileName:
            asd['fileName'] = fileName
        elif requisites:
            asd['requisites'] = requisites
        else:
            asd['fImageId'] = data['fImageId']
            asd['sImageId'] = data['sImageId']
        asd['username'] = msg.from_user.username
        asd['userid'] = msg.from_user.id
        await opponentState.update_data(asd)
    
    for i in checker_ids:
        await msg.bot.send_message(i, f'{text1}\n{text2}', reply_markup=keyboard, parse_mode='MarkdownV2')
    if isinstance(msg, Message):
        await msg.answer(text='Отправлено на проверку админу @dokuzu и @crepecafe')
    else:
        await msg.message.answer(text='Отправлено на проверку админу @dokuzu и @crepecafe')


async def ReportToAdmin(msg, fields, photoId):
    text=f'Попытка сканирования @{msg.from_user.username} \n'
    if fields:
        for i in fields:
            text+=f'{i}-{fields[i] if fields[i] is not None else "нет данных"} \n'
    else:
        text = 'ошибка при сканировании'
    for i in checker_ids:
        await msg.bot.send_photo(i, photo=photoId, disable_notification=True)
        await msg.bot.send_message(i, text, disable_notification=True)


def CreateDataFile(data, username):
    with open(f'{PATHS["data"]}{username}.data', 'w') as f:
        f.write(f'iin={data["ИИН"]}||email=desftrolks@mail.ru')